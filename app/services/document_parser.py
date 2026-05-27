import logging
import os
import uuid
from pathlib import Path
from typing import List, Tuple

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from app.core.config import settings

logger = logging.getLogger(__name__)


class UnsupportedFileType(Exception):
    pass


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    return "\n".join(pages)


def _read_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


_READERS = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".txt": _read_text,
    ".md": _read_text,
}


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    reader = _READERS.get(suffix)
    if reader is None:
        raise UnsupportedFileType(f"unsupported file type: {suffix}")
    return reader(path)


def chunk_text(text: str, source: str, doc_id: str) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_text(text)
    docs: List[Document] = []
    for i, chunk in enumerate(chunks):
        docs.append(
            Document(
                page_content=chunk,
                metadata={"source": source, "doc_id": doc_id, "chunk_index": i},
            )
        )
    return docs


def save_upload(filename: str, content: bytes) -> Tuple[str, Path]:
    """Persist uploaded bytes and return (doc_id, path)."""
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    doc_id = uuid.uuid4().hex
    safe_name = Path(filename).name
    target = Path(settings.UPLOAD_DIR) / f"{doc_id}_{safe_name}"
    target.write_bytes(content)
    return doc_id, target


def parse_and_chunk(filename: str, content: bytes) -> Tuple[str, List[Document]]:
    """End-to-end: persist, extract, chunk. Returns (doc_id, chunks)."""
    doc_id, path = save_upload(filename, content)
    text = extract_text(path)
    if not text.strip():
        logger.warning("Extracted empty text from %s", filename)
    chunks = chunk_text(text, source=filename, doc_id=doc_id)
    logger.info("Parsed %s into %d chunks (doc_id=%s)", filename, len(chunks), doc_id)
    return doc_id, chunks
